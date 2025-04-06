import os
import logging
from typing import List, Dict, Any
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_RIGHT, TA_LEFT, TA_CENTER, TA_JUSTIFY

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants for formatting
PAGE_WIDTH = 6  # inches
PAGE_HEIGHT = 9  # inches
# Use standard ReportLab fonts instead of Georgia
FONT_NAME = 'Times-Roman'  # Standard PDF font that's always available
DOCX_FONT_NAME = 'Georgia'  # Keep Georgia for Word docs
BODY_FONT_SIZE = 11
CHAPTER_TITLE_FONT_SIZE = 30
DROPCAP_SIZE = 30
CHARS_PER_PAGE = 3000  # Approximate


def divide_text_into_pages(text: str, max_pages: int = 10, chars_per_page: int = CHARS_PER_PAGE) -> List[str]:
    """
    Divide text into pages based on approximate character count.
    
    Args:
        text: The text to divide
        max_pages: Maximum number of pages to return
        chars_per_page: Approximate characters per page
        
    Returns:
        List of page content strings
    """
    pages = []
    
    # Simple division by character count
    for i in range(0, len(text), chars_per_page):
        pages.append(text[i:i+chars_per_page])
        if len(pages) >= max_pages:
            break
    
    return pages


def apply_drop_cap(paragraph):
    """
    Apply drop cap effect to the first letter of a paragraph.
    
    Args:
        paragraph: docx paragraph object
    """
    if not paragraph.runs:
        return
    
    first_run = paragraph.runs[0]
    if not first_run.text:
        return
    
    # Extract first letter and rest of text
    first_letter = first_run.text[0]
    rest_of_text = first_run.text[1:]
    
    # Clear existing text and format with drop cap
    first_run.clear()
    
    # Add styled first letter (drop cap)
    drop_cap_run = paragraph.add_run(first_letter)
    drop_cap_run.font.size = Pt(DROPCAP_SIZE)
    drop_cap_run.font.name = DOCX_FONT_NAME
    
    # Add rest of text
    normal_run = paragraph.add_run(rest_of_text)
    normal_run.font.size = Pt(BODY_FONT_SIZE)
    normal_run.font.name = DOCX_FONT_NAME


def add_page_number(document, page_number: int):
    """
    Add page number with proper alignment based on odd/even page.
    
    Args:
        document: docx document object
        page_number: Current page number
    """
    paragraph = document.add_paragraph()
    paragraph.text = f"{page_number}"
    
    # Left align for even pages, right align for odd pages
    if page_number % 2 == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def create_book_document(translated_chapters: List[Dict[str, Any]], output_path: str = "livro_traduzido.docx") -> str:
    """
    Create a formatted book document from translated chapters.
    
    Args:
        translated_chapters: List of dictionaries with 'id' and 'content' keys
        output_path: Path to save the output document
        
    Returns:
        Path to the created document
    """
    # Sort chapters by ID to ensure correct order regardless of input order
    sorted_chapters = sorted(translated_chapters, key=lambda x: x["id"])
    
    # Enhanced logging to diagnose multi-chapter issues
    logger.info(f"Starting DOCX creation for {len(sorted_chapters)} chapters at {output_path}")
    
    # Log chapter IDs being processed
    chapter_ids = [chapter["id"] for chapter in sorted_chapters]
    logger.info(f"Chapter IDs to be included (sorted): {chapter_ids}")
    
    document = Document()
    
    # Set page size to 6 x 9 inches
    section = document.sections[0]
    section.page_width = Inches(PAGE_WIDTH)
    section.page_height = Inches(PAGE_HEIGHT)
    
    # Configure margins
    section.left_margin = Inches(0.75)  # Wider margin near binding
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    page_number = 1
    
    # Process each chapter
    for i, chapter in enumerate(sorted_chapters):
        chapter_id = chapter["id"]
        chapter_content = chapter["content"]
        
        # Detailed logging for each chapter
        logger.info(f"Processing chapter {i+1}/{len(sorted_chapters)}: ID={chapter_id}, Content length={len(chapter_content)} chars")
        
        # Add chapter title
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(f"Capítulo {chapter_id}")
        title_run.font.size = Pt(CHAPTER_TITLE_FONT_SIZE)
        title_run.font.name = DOCX_FONT_NAME  # Use Georgia for Word docs
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Divide chapter into pages
        pages = divide_text_into_pages(chapter_content)
        logger.info(f"Chapter {chapter_id} divided into {len(pages)} pages")
        
        for j, page_content in enumerate(pages):
            content_paragraph = document.add_paragraph()
            content_run = content_paragraph.add_run(page_content)
            content_run.font.size = Pt(BODY_FONT_SIZE)
            content_run.font.name = DOCX_FONT_NAME  # Use Georgia for Word docs
            
            # Apply drop cap to first paragraph of the chapter
            if j == 0:
                apply_drop_cap(content_paragraph)
            
            # Add page number
            add_page_number(document, page_number)
            page_number += 1
            
            # Add page break if not the last page of the book
            if j < len(pages) - 1 or i < len(sorted_chapters) - 1:
                logger.info(f"Adding page break after page {j+1} of chapter {chapter_id}")
                document.add_page_break()
    
    # Create the output directory if it doesn't exist
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Save the document
    document.save(output_path)
    logger.info(f"Book document created successfully: {output_path}")
    
    return output_path


def create_pdf_document(translated_chapters: List[Dict[str, Any]], output_path: str = "livro_traduzido.pdf") -> str:
    """
    Create a formatted PDF document from translated chapters using ReportLab.
    
    Args:
        translated_chapters: List of dictionaries with 'id' and 'content' keys
        output_path: Path to save the output PDF
        
    Returns:
        Path to the created PDF
    """
    # Sort chapters by ID to ensure correct order regardless of input order
    sorted_chapters = sorted(translated_chapters, key=lambda x: x["id"])
    
    # Enhanced logging to diagnose multi-chapter issues
    logger.info(f"Starting PDF creation for {len(sorted_chapters)} chapters at {output_path}")
    
    # Log chapter IDs being processed
    chapter_ids = [chapter["id"] for chapter in sorted_chapters]
    logger.info(f"Chapter IDs to be included (sorted): {chapter_ids}")
    
    # Create output directory if needed
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Create PDF document
    doc = SimpleDocTemplate(
        output_path,
        pagesize=(PAGE_WIDTH*inch, PAGE_HEIGHT*inch),
        leftMargin=0.75*inch,
        rightMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Create custom styles with standard fonts
    try:
        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontName='Times-Bold',  # Standard PDF font
            fontSize=CHAPTER_TITLE_FONT_SIZE,
            alignment=TA_RIGHT,  # Right alignment
            spaceAfter=24
        )
        
        body_style = ParagraphStyle(
            'BodyText',
            parent=styles['Normal'],
            fontName='Times-Roman',  # Standard PDF font
            fontSize=BODY_FONT_SIZE,
            leading=14,  # Line spacing
            firstLineIndent=20,  # First line indentation
            alignment=TA_JUSTIFY  # Justified text
        )
        
        logger.info("Using standard PDF fonts for document generation")
    except Exception as e:
        logger.error(f"Error creating paragraph styles: {str(e)}")
        # Fallback to basic styles if custom styles fail
        chapter_title_style = styles['Heading1']
        body_style = styles['Normal']
    
    # Build document content
    story = []
    
    # Import PageBreak here to ensure it's available
    from reportlab.platypus import PageBreak
    
    for i, chapter in enumerate(sorted_chapters):
        chapter_id = chapter["id"]
        chapter_content = chapter["content"]
        
        # Detailed logging for each chapter
        logger.info(f"Processing chapter {i+1}/{len(sorted_chapters)}: ID={chapter_id}, Content length={len(chapter_content)} chars")
        
        # Add chapter title
        try:
            # Start each chapter on a new page (except the first one)
            if i > 0:
                logger.info(f"Adding page break before chapter {chapter_id}")
                story.append(PageBreak())
            
            # Add the chapter title
            title = Paragraph(f"<para>Capítulo {chapter_id}</para>", chapter_title_style)
            story.append(title)
            story.append(Spacer(1, 0.2*inch))
            
            # Add chapter content with paragraphs
            paragraphs = chapter_content.split('\n\n')
            logger.info(f"Chapter {chapter_id} split into {len(paragraphs)} paragraphs")
            
            for j, p in enumerate(paragraphs):
                if p.strip():
                    # Escape any XML characters that might cause issues
                    p = p.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    para = Paragraph(p, body_style)
                    story.append(para)
                    story.append(Spacer(1, 0.1*inch))
            
        except Exception as e:
            logger.error(f"Error processing chapter {chapter_id}: {str(e)}")
            # Continue processing other chapters even if one fails
    
    # Log the final story length
    logger.info(f"PDF story built with {len(story)} elements")
    
    # Build the PDF
    try:
        doc.build(story)
        logger.info(f"PDF document created successfully: {output_path}")
    except Exception as e:
        logger.error(f"Error building PDF document: {str(e)}")
        raise
    
    return output_path 